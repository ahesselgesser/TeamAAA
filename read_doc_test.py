import docx
from docx.oxml.ns import qn
import re

def regex_inc(regex_list, regex_counter, match_list):
    regex = re.findall(regex_list[regex_counter], para.text)
    if (regex):
        #print(regex[0])
        match_list.append(regex[0])
        if (regex_counter < len(regex_list) - 1):
            regex_counter += 1
    return (regex_counter, match_list)


## Change this file location to the location of the file you are parsing
doc = docx.Document("C:/Users/twins/Desktop/UNO classes/Spring 2021 Semester/CSCI 4970 - Capstone/Python tests/undergrad2018-regular.docx")

# 'C:\Users\twins\Desktop\UNO classes\Spring 2021 Semester\CSCI 4970 - Capstone\Python tests'
# https://github.com/ahesselgesser/TeamAAA

#### Variables to capture
## College, Department/School ?, Program, Degree Level, Academic Year of Report,
## Date Range of Reported Data, Person Preparing the Report, SLOs (maybe in an array, or tuple, or something)
## Bloom's Taxonomy (checkboxes)

### Creating lists to hold all the different regex's and matches, then iterate through them.
regex_counter = 0
regex_list = ['College:\s*(.*)\s*Department/School:', 'Department/School:\s*(.*)', 'Program:\s*(.*)\s*Degree Level:', 'Degree Level:\s*(.*)', 'Academic Year of Report:\s*(.*)\s*Date', 'Data:\s*(.*)']

### Match 1 is Colleges
### Match 2 is Department/School
### 
match_list = []

all_paras = doc.paragraphs
for para in all_paras:
    (regex_counter, match_list) = regex_inc(regex_list, regex_counter, match_list)
    (regex_counter, match_list) = regex_inc(regex_list, regex_counter, match_list)

print(match_list)


"""
### https://stackoverflow.com/questions/27861732/parsing-of-table-from-docx-file/27862205 ###
data = []

table = doc.tables[0]
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)
    #print(row.cells)
    
    # Establish the mapping based on the first row
    # headers; these will become the keys of our dictionary
    if i == 0:
        keys = tuple(text)
        continue

    # Construct a dictionary for this row, mapping
    # keys to values for this row
    row_data = dict(zip(keys, text))
    data.append(row_data)

print(data)
###    https://stackoverflow.com/questions/27861732/parsing-of-table-from-docx-file/27862205 ###
"""