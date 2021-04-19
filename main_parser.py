import docx
from docx.oxml.ns import qn
import re
import read_doc_chx


def regex_inc(regex_list, regex_counter, match_list, para):
    if(regex_counter < len(regex_list)):
        regex = re.findall(regex_list[regex_counter], para.text)
        if (regex):
            match_list.append(regex[0])
            regex_counter += 1
    return (regex_counter, match_list)

def table_two_access(file_name, column, table, row_iter, data, one_cols, max_rows):
    if (file_name == "undergrad2018-regularv2.docx" or file_name == "undergrad2019-regular.docx"):
        for row in table.rows:
            if row_iter <= one_cols or (row_iter > max_rows and row_iter <= one_cols + max_rows):
                data.append({"Row" + str(row_iter): row.cells[column[1]].text})
            if (row_iter > one_cols and row_iter < max_rows + 1) or (row_iter > max_rows + one_cols):
                data.append({row.cells[0].text: row.cells[column[2]].text})
            row_iter += 1
            if row_iter > max_rows:
                row_iter = 0
    else:
        for row in table.rows:
            if (row_iter == 0):
                row_iter += 1
                continue
            data.append(("SLO " + row.cells[column[1]].text, row.cells[column[2]].text, row.cells[column[3]].text, row.cells[column[4]].text))
    return data

def table_three_access(column, table, row_iter, data, one_cols, max_rows, slo_count):
    ### Original loop from link below
    ### https://www.reddit.com/r/learnpython/comments/dbie6s/help_iterating_over_a_table_in_pythondocx/
    for row in table.rows:
        if row.cells[column[1]].text == "SLO " + str(slo_count + 1) + ": ":
            break
        if row_iter <= one_cols or (row_iter > max_rows and row_iter <= one_cols + max_rows):
            data.append({"Row" + str(row_iter): row.cells[column[1]].text})
        if (row_iter > one_cols and row_iter < max_rows + 1) or (row_iter > max_rows + one_cols):
            data.append({row.cells[0].text: row.cells[column[2]].text})
        row_iter += 1
        if row_iter > max_rows:
            row_iter = 0
    return data

def table_four_access(file_name, column, table, row_iter, data_coll_list, slo_count):
    if (file_name == "undergrad2018-regularv2.docx" or file_name == "undergrad2019-regular.docx"):
        for row in table.rows:
            if "SLO " + str(slo_count + 1) in row.cells[column[1]].text:
                break
            slo_measure = row.cells[column[1]].text
            data_coll_range = row.cells[column[2]].text
            num_students = row.cells[column[3]].text
            percent_students = row.cells[column[4]].text
            extend_values = [(slo_measure, data_coll_range, num_students, percent_students)]
            if data_coll_range == "" and num_students == "" and percent_students == "":
                continue
            data_coll_list.extend(extend_values)
    else:
        for row in table.rows:
            if "SLO " + str(slo_count + 1) in row.cells[column[1]].text:
                break
            slo_num = row.cells[column[1]].text
            dec_and_act = row.cells[column[2]].text
            extend_values = [(slo_num, dec_and_act)]
            data_coll_list.extend(extend_values)
    return data_coll_list
def run():
    print("Choose a file_name:")
    print("1: undergrad2018-regular.docx")
    print("2: undergrad2019-regular.docx")
    print("3: undergrad2019-accredited.docx")
    print("4: grad-accredited.docx")

    file_choice = input()
    incorrect_choice = True

    while (incorrect_choice):
        if (file_choice == "1"):
            file_name = "undergrad2018-regularv2.docx"
        elif (file_choice == "2"):
            file_name = "undergrad2019-regular.docx"
        elif (file_choice == "3"):
            file_name = "undergrad2019-accredited.docx"
        elif (file_choice == "4"):
            file_name = "grad-accredited.docx"
        else:
            file_choice = input("Incorrect choice, try again: ")
            continue
        break

    unzip_name = file_name[:-5] + ".zip"
    zip_dir = file_name[:-5] + "/"
    file_dir = "C:/Users/twins/Desktop/UNO classes/Spring 2021 Semester/CSCI 4970 - Capstone/Python tests/"
    table_access_list = ()

    if (file_name == "undergrad2018-regularv2.docx" or file_name == "undergrad2019-regular.docx"):
        regex_header_list = ['College:\s*(.*)\s*Department/School:', 'Department/School:\s*(.*)', 'Program:\s*(.*)\s*Degree Level:', 'Degree Level:\s*(.*)', 'Academic Year of Report:\s*(.*)\s*Date', 'Data:\s*(.*)',
            'Person Preparing the Report:\s(.*)']
        if (file_name == "undergrad2018-regularv2.docx"):
            table_access_list = (0, 1, 2, 3, 5)
        if (file_name == "undergrad2019-regular.docx"):
            table_access_list = (0, 1, 2, 4, 6)
    else:
        regex_header_list = ['College:\s*(.*)\s*Department/School:', 'Department/School:\s*(.*)', 'Program:\s*(.*)\s*Degree Level:', 'Degree Level:\s*(.*)', 'Academic Year of Report:\s*(.*)\s*Person', 
            'Person Preparing the Report:\s(.*)', 'Last Accreditation Review:\s(.*)\s*Accreditation', 'Accreditation Body:\s(.*)\s*']   
        table_access_list = (0, 1, 2, 3)

    read_doc_chx.copy_and_unzip(file_dir, file_name, unzip_name, zip_dir)

    xml_string = "word/document.xml"
    xml_path = file_dir + zip_dir + xml_string

    (chkbox_element_list, slo_count) = read_doc_chx.find_checkbox_elements(xml_path)
    print(chkbox_element_list)
    print(slo_count)

    ## Change this file location to the location of the file you are parsing
    doc = docx.Document(file_dir + file_name)

    ### Creating lists to hold all the different regex's and matches, then iterate through them.
    regex_counter = 0
    
    match_list = []

    all_paras = doc.paragraphs

    for para in all_paras:
        (regex_counter, match_list) = regex_inc(regex_header_list, regex_counter, match_list)
        (regex_counter, match_list) = regex_inc(regex_header_list, regex_counter, match_list)
    print(match_list)

    data = []

    slo_list = []

    ### Original code from https://stackoverflow.com/questions/27861732/parsing-of-table-from-docx-file/27862205 ###
    #### Append the header rows as keys, and the cells beneath them as the corresponding values
    #### This is for the SLO table
    table = doc.tables[table_access_list[0]]
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        
        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue
        elif i == 2 * (slo_count - 1):
            break

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
    ## Append the SLOs to the SLO list
    for slo in range (0, slo_count):
        slo_list.append(data[slo]['Student Learning Outcomes'])

    data.clear()

    print(slo_list)

    column = (0, 0, 1, 2, 3)
    row_iter = 0
    one_cols = 2
    max_rows = 10
    table = doc.tables[table_access_list[1]]
    prev_text = ""
    data = table_two_access(file_name, column, table, row_iter, data)


    row_iter = 0
    prev_text = ""
    table = doc.tables[table_access_list[2]]
    data = table_three_access(column, table, row_iter, data)

    print("\nData:")
    print(data)

    #### DATA COLLECTION AND ANALYSIS
    data_coll_list = []
    table = doc.tables[table_access_list[3]]
    data_coll_list = table_four_access(file_name, column, table, row_iter, data_coll_list)

    print("\nData collection list:")
    print(data_coll_list)

    #### DECISIONS AND ACTIONS
    if (file_name == "undergrad2018-regularv2.docx" or file_name == "undergrad2019-regular.docx"):
        table = doc.tables[table_access_list[4]]
        dec_act = []
        for row in table.rows:
            slo_num = row.cells[column[1]].text
            slo_act = row.cells[column[2]].text
            if slo_act == "\n\n\n\n" or slo_act == "":
                continue
            dec_act.extend([(slo_num, slo_act)])

        print(dec_act)