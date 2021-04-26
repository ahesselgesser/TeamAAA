import docx
from docx.oxml.ns import qn
import re
from mysite.core.paser import read_doc_chx
import psycopg2
from mysite.core.paser import insertData

def assessmentMethodProcessor(assessment_methods):
    output = []
    temp = []
    for item in assessment_methods:
        if (type(item) == str and item.startswith("SLO")):
            if (len(temp) != 0):
                output.append(temp)
                temp = []
        temp.append(item)
    output.append(temp)
    return output

def regex_inc(regex_list, regex_counter, match_list, para):
    """!  Allows a for loop to loop through a list of regular expressions.
    
    @param regex_list - This is the list of regular expressions the loop accesses.
    @param regex_counter - This counter keeps track of how many times the regular expression list has been accessed.
    @param match_list - This list keeps track of all the matches that the regular expressions have caught.
    @param para - This is the text that the regular expression is compared to.

    @return The counter keeping track of regular expressions and the list of matches.
    """
    if(regex_counter < len(regex_list)):
        regex = re.findall(regex_list[regex_counter], para.text)
        if (regex):
            match_list.append(regex[0])
            regex_counter += 1
    return (regex_counter, match_list)

def table_two_access(is_undergrad, column, table, row_iter, data, one_cols, max_rows):
    """!  Captures the information in the second table of the document.
    
    @param is_undergrad - Boolean value, True if the report is for an undergraduate degree program.
    @param column - A tuple that refers to the column number we are accessing in the table.
    @param table - The docx.Table object that we are accessing.
    @param row_iter - An integer that keeps track of what row we are accessing.
    @param data - A list that the captured data is appended to.
    @param one_cols - I honestly forgot how I was using this variable.
    @param max_rows - The maximum number of rows that should be in the table we are accessing.

    @return The captured data in a list.
    """    
    if (is_undergrad):
        for row in table.rows:
            if row_iter <= one_cols or (row_iter > max_rows and row_iter <= one_cols + max_rows):
                data.append(row.cells[column[1]].text)
            if (row_iter > one_cols and row_iter < max_rows + 1) or (row_iter > max_rows + one_cols):
                data.append((row.cells[0].text, row.cells[column[2]].text))
            row_iter += 1
            if row_iter > max_rows:
                row_iter = 0
    else:
        for row in table.rows:
            if (row_iter == 0):
                row_iter += 1
                continue
            data.append(("\n\n" + "SLO " + row.cells[column[1]].text, row.cells[column[2]].text, row.cells[column[3]].text, row.cells[column[4]].text))
    return data

def table_three_access(column, table, row_iter, data, one_cols, max_rows, slo_count):
    """!  Captures the information in the third table of the document.
    
    @param column - A tuple that refers to the column number we are accessing in the table.
    @param table - The docx.Table object that we are accessing.
    @param row_iter - An integer that keeps track of what row we are accessing.
    @param data - A list that the captured data is appended to.
    @param one_cols - I honestly forgot how I was using this variable.
    @param max_rows - The maximum number of rows that should be in the table we are accessing.
    @param slo_count - This integer is the number of SLOs that were counted in the document.

    @return The captured data in a list.
    """    

    ### Original loop from link below
    ### https://www.reddit.com/r/learnpython/comments/dbie6s/help_iterating_over_a_table_in_pythondocx/
    for row in table.rows:
        if row.cells[column[0]].text == "SLO 1":
            break
        if row_iter <= one_cols or (row_iter > max_rows and row_iter <= one_cols + max_rows):
            data.append(row.cells[column[1]].text)
        if (row_iter > one_cols and row_iter < max_rows + 1) or (row_iter > max_rows + one_cols):
            data.append((row.cells[0].text, row.cells[column[2]].text))
        row_iter += 1
        if row_iter > max_rows:
            row_iter = 0
    return data

def table_four_access(is_undergrad, column, table, row_iter, data_coll_list, slo_count):
    """!  Captures the information in the fourth table of the document.
    
    @param is_undergrad - Boolean value, True if the report is for an undergraduate degree program.
    @param column - A tuple that refers to the column number we are accessing in the table.
    @param table - The docx.Table object that we are accessing.
    @param row_iter - An integer that keeps track of what row we are accessing.
    @param data_coll_list - A list that the captured data is appended to.
    @param one_cols - I honestly forgot how I was using this variable.
    @param slo_count - This integer is the number of SLOs that were counted in the document.

    @return The captured data in a list of tuples.
    """
            
    if (is_undergrad):
        for row in table.rows:
            if "SLO " + str(slo_count + 1) in row.cells[column[1]].text:
                break
            slo_measure = row.cells[column[1]].text
            data_coll_range = row.cells[column[2]].text
            num_students = row.cells[column[3]].text
            percent_students = row.cells[column[4]].text
            extend_values = [(slo_measure, data_coll_range, num_students, percent_students)]
            if data_coll_range == "" and num_students == "" and percent_students == "":
                continue
            data_coll_list.extend(extend_values)
    else:
        for row in table.rows:
            slo_num = row.cells[column[1]].text
            dec_and_act = row.cells[column[2]].text
            extend_values = [(slo_num, dec_and_act)]
            data_coll_list.extend(extend_values)
            row_iter += 1
    return data_coll_list

def run(uploaded_filename):
    """! This function captures the information from documents.

    @param uploaded_filename This is the file name of the document we are running the method on.
    """

    ## Whether the report is undergraduate; False = graduate, True = undergraduate.
    is_undergrad = False
    ## Whether the report is accredited; False = Non-accredited, True = accredited 
    is_accredited = False
    ## The year the report was created; 18 = 2018, 19 = 2019, Any other numbers should not be allowed.
    report_year = 0

    ## The regular expression used to capture the "B" or "M" for Bachelors or Masters, in the file name. "B" = undergraduate, anything else = graduate.
    filename_regex = "\.*?\((.{1})\w+\)\.*"
    ## The regular expression used to capture if the report was from 2019. Returns a regular expression match object.
    report_year_regex = 'Academic Year of Report:\s*(2019)\s*Date'

    ## The match object returned when checking the uploaded file name against the file name regular expression defined above.
    filename_match = re.search(filename_regex, uploaded_filename)

    if (filename_match.group(1).lower() == "b"):
        is_undergrad = True

    ## The name we use to create the ZIP version of the document. It should be "filename".zip.
    unzip_name = uploaded_filename[:-5] + ".zip"
    ## The directory that we use to store the documents after unzipping the ZIP file.
    zip_dir = uploaded_filename[:-5] + "/"
    ## The file directory that we store all the documents.
    file_dir = "./media/"
    ## A tuple that stores the table numbers where the document information is stored.
    table_access_list = ()

    ## The docx object that we use to access the document information gained from the docx module.
    doc = docx.Document(file_dir + uploaded_filename)

    ## Checking to see if the report year of the document is 2019.
    for para in doc.paragraphs:
        report_year_match = re.search(report_year_regex, para.text)
        if(report_year_match != "None"):
            report_year = 19
            break

    ## Checking to see if the report is for an Accredited degree. Otherwise assigns the report year.
    if("Accredited" in doc.paragraphs[0].text):
        is_accredited = True
    elif(report_year == 0 and "19" in doc.paragraphs[1].text):
        report_year = 19
    elif(report_year == 0):
        report_year = 18

    read_doc_chx.copy_and_unzip(file_dir, uploaded_filename, unzip_name, zip_dir)

    ## Determines the appropriate regular expression list for the document header and table access based on undergraduate/graduate, accredited/non, and report year.
    if (is_undergrad and not is_accredited ):
        regex_header_list = ['College:\s*(.*)\s*Department/School:', 'Department/School:\s*(.*)', 'Program:\s*(.*)\s*Degree Level:', 'Degree Level:\s*(.*)', 'Academic Year of Report:\s*(.*)\s*Date', 'Date Range of Reported Data:\s*(.*)',
            'Person Preparing the Report:\s(.*)']
        report_header_list = ['College: ', 'Department/School: ', 'Program: ', 'Degree Level: ', 'Academic Year of Report: ', 'Date: ', 'Person Preparing the Report: ']
        if (report_year == 18):
            table_access_list = (0, 1, 2, 3, 5)
        if (report_year == 19):
            table_access_list = (0, 1, 2, 4, 6)
    else:
        regex_header_list = ['College:\s*(.*)\s*Department/School:', 'Department/School:\s*(.*)', 'Program:\s*(.*)\s*Degree Level:', 'Degree Level:\s*(.*)', 'Academic Year of Report:\s*(.*)\s*Person', 
            'Person Preparing the Report:\s(.*)', 'Last Accreditation Review:\s(.*)\s*Accreditation', 'Accreditation Body:\s(.*)\s*']   
        report_header_list = ['College: ', 'Department/School: ', 'Program: ', 'Degree Level: ', 'Academic Year of Report', 'Person Preparing the Report: ', 'Last Accreditation Review: ', 'Accreditation Body: ']
        table_access_list = (0, 1, 2, 3)

    ## The path in the unzipped directory to the document XML file.
    xml_string = "word/document.xml"
    ## The full path to the unzipped directory.
    xml_path = file_dir + zip_dir + xml_string

    #assessment_obj = assessment_models.Assessment()
    
    ## Store the return values from the find_checkbox_elements function.
    (chkbox_element_list, slo_count, list_of_lists) = read_doc_chx.find_checkbox_elements(xml_path)

    ## Printing out the checkbox elements, ONLY USED FOR TESTING PURPOSES.
    print("Checkbox elements here: printing at Line 209")
    print("============================================================================================")

    for item in list_of_lists:
        print(item)
    print("============================================================================================\n")

    ## A counter to keep track of what regular expression in a list we are on.
    regex_counter = 0
    
    ## A list object used to hold the matches found from the regular expressions.
    match_list = []

    ## Stores all the docx Paragraph objects found in the document.
    all_paras = doc.paragraphs

    ## Loop through all of the paragraphs and compare them to the regular expressions.
    for para in all_paras:
        (regex_counter, match_list) = regex_inc(regex_header_list, regex_counter, match_list, para)
        (regex_counter, match_list) = regex_inc(regex_header_list, regex_counter, match_list, para)

    ## Print out the report header info, ONLY FOR TESTING PURPOSES.
    print("Report Header info: printing at line 231")
    print("============================================================================================")
    
    for i in range(0, len(match_list)):
        print(report_header_list[i] + ": " + match_list[i])
    print("============================================================================================\n")

    ## List used to store information captured from tables.
    data = []

    ## List used to store SLOs from document.
    slo_list = []

    ### Original code from https://stackoverflow.com/questions/27861732/parsing-of-table-from-docx-file/27862205 ###
    #### This is for the SLO table

    ## Used to store the first docx Table object that we are interested in.
    table = doc.tables[table_access_list[0]]
    ## Append the header rows as keys, and the cells beneath them as the corresponding values
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        
        ## Establish the mapping based on the first row headers; these will become the keys of our dictionary.
        if i == 0:
            keys = tuple(text)
            continue
        elif i == 2 * (slo_count - 1):
            break

        ## Construct a dictionary for this row, mapping keys to values for this row.
        row_data = dict(zip(keys, text))
        data.append(row_data)
        
    ## Append the SLOs to the SLO list
    for slo in range (0, slo_count):
        slo_list.append(data[slo]['Student Learning Outcomes'])

    ## Clear data so that we can use it again.
    data.clear()

    ## Print out the list of SLOs, ONLY FOR TESTING PURPOSES.
    print("SLO List: printing at line 272")
    print("============================================================================================")
    for slo in slo_list:
        print(slo)
    print("============================================================================================\n")

    ## Tuple to hold the 0-based number for each column in a table.
    column = (0, 0, 1, 2, 3)
    ## Integer used to keep track of what row we are accessing on a table.
    row_iter = 0
    ## I still don't remember how I created the logic for this, sorry.
    one_cols = 2
    ## The max rows that should be in the tables we are accessing. The XML for these tables is shot to hell, so we had to create a variable to tell the loop to end.
    max_rows = 10
    ## Set the table variable to the next table we need to access, via the table_access_list.
    table = doc.tables[table_access_list[1]]
    ## Reset the previous text. Maybe obsolete.
    prev_text = ""
    data = table_two_access(is_undergrad, column, table, row_iter, data, one_cols, max_rows)

    ## Reset row_iter, this might not be needed if the row_iter variable is not changed after calling table_two_access(...).
    row_iter = 0
    ## Maybe obsolete.
    prev_text = ""
    ## Set the table to the next table in the table_access_list.
    table = doc.tables[table_access_list[2]]
    assessment_methods = table_three_access(column, table, row_iter, data, one_cols, max_rows, slo_count)

    #added in splitting assessment_methods into one list per SLO
    assessment_methods = assessmentMethodProcessor(assessment_methods)

    ## Print out the Assessment Methods info, ONLY FOR TESTING PURPOSES.
    print("Assessment Methods info: printing at line 211")
    print("============================================================================================")
    for method_info in assessment_methods:
        print(method_info)
    print("============================================================================================\n")

    #### DATA COLLECTION AND ANALYSIS
    ## The list that holds the information from the Data Collection & Analysis table.
    data_coll_list = []
    ## Set the table variable to the next table.
    table = doc.tables[table_access_list[3]]
    data_coll_list = table_four_access(is_undergrad, column, table, row_iter, data_coll_list, slo_count)

    ## Print out the Data Collection & Analysis info, ONLY FOR TESTING purposes.
    skip = 0
    print("Data Collection & Analysis info: printing at line 223")
    print("============================================================================================")
    if (is_accredited):
        for data_coll in data_coll_list:
            print(data_coll[0] + " : " + data_coll[1])
    else:
        for data_coll in data_coll_list:
            if skip == 0:
                skip = 1
                continue
            print(data_coll)
    print("============================================================================================\n")

    #### DECISIONS AND ACTIONS
    ## This is only for undergraduate reports. 
    if (is_undergrad):
        table = doc.tables[table_access_list[4]]
        dec_act = []
        for row in table.rows:
            slo_num = row.cells[column[1]].text
            slo_act = row.cells[column[2]].text
            if slo_act == "\n\n\n\n" or slo_act == "":
                continue
            dec_act.extend([(slo_num, slo_act)])

        ## Print out Decisions & Actions info, ONLY FOR TESTING PURPOSES.
        print("Decisions & Actions info: printing at line 248")
        print("============================================================================================")
        for info in dec_act:
            print(info[0] + info[1])
        print("============================================================================================\n")
       # insertData.insertCheckBox(list_of_lists)
       # print(len(slo_list))
        insertData.insertReportHeader(match_list, list_of_lists,is_accredited,len(slo_list), dec_act, assessment_methods, slo_list, data_coll_list)