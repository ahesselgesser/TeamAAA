import docx
from docx.oxml.ns import qn
import re
import read_doc_chx

def regex_inc(regex_list, regex_counter, match_list):
    if(regex_counter < len(regex_list)):
        regex = re.findall(regex_list[regex_counter], para.text)
        if (regex):
            match_list.append(regex[0])
            #if (regex_counter < len(regex_list) - 1):
            regex_counter += 1
    return (regex_counter, match_list)

accredited_file_dir = "C:/Users/twins/Desktop/UNO classes/Spring 2021 Semester/CSCI 4970 - Capstone/Python tests/"
accredited_file_name = "undergrad2019-accredited.docx"

#read_doc_chx.copy_and_unzip(accredited_file_name)
(chkbox_element_list, slo_count) = read_doc_chx.find_checkbox_elements()
print(slo_count)

## Change this file location to the location of the file you are parsing
doc = docx.Document(accredited_file_dir + accredited_file_name)

# 'C:\Users\twins\Desktop\UNO classes\Spring 2021 Semester\CSCI 4970 - Capstone\Python tests'
# https://github.com/ahesselgesser/TeamAAA

#### Variables to capture
## College, Department/School ?, Program, Degree Level, Academic Year of Report,
## Date Range of Reported Data, Person Preparing the Report, SLOs (maybe in an array, or tuple, or something)
## Bloom's Taxonomy (checkboxes)

### Creating lists to hold all the different regex's and matches, then iterate through them.
regex_counter = 0
ug18_regex_header_list = ['College:\s*(.*)\s*Department/School:', 'Department/School:\s*(.*)', 'Program:\s*(.*)\s*Degree Level:', 'Degree Level:\s*(.*)', 'Academic Year of Report:\s*(.*)\s*Person', 
'Person Preparing the Report:\s(.*)', 'Last Accreditation Review:\s(.*)\s*Accreditation', 'Accreditation Body:\s(.*)\s*']   

### Match 1 is Colleges
### Match 2 is Department/School
### 
match_list = []

all_paras = doc.paragraphs

for para in all_paras:
    (regex_counter, match_list) = regex_inc(ug18_regex_header_list, regex_counter, match_list)
    (regex_counter, match_list) = regex_inc(ug18_regex_header_list, regex_counter, match_list)
    #if (regex_counter >= 7):
        #print(para.text)
print(match_list)

data = []

### UG 2019 Regular Tables
## SLO Table
# Checkboxes
## SLO communication to stakeholders
## Assessment Methods
# Separate table for EACH SLO
# Will have to devise a method to determine when these tables end and the next category begins
# Checkboxes in these tables
## Data Collection And Analysis Table
# Two tables
# 
## Resuls communicated within program Table
## Decisions & Actions Table
### END
slo_list = []

### Original code from https://stackoverflow.com/questions/27861732/parsing-of-table-from-docx-file/27862205 ###
#### Append the header rows as keys, and the cells beneath them as the corresponding values
#### This is for the SLO table
table = doc.tables[0]
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)
    
    # Establish the mapping based on the first row
    # headers; these will become the keys of our dictionary
    if i == 0:
        keys = tuple(text)
        continue
    elif i == 2 * (slo_count - 1):
        break

    # Construct a dictionary for this row, mapping
    # keys to values for this row
    row_data = dict(zip(keys, text))
    data.append(row_data)
## Append the SLOs to the SLO list
for slo in range (0, slo_count):
    slo_list.append(data[slo]['Student Learning Outcomes'])

data.clear()

print(slo_list)

first_col = 0
sec_col = 1
third_col = 2
fourth_col = 3
row_iter = 0
one_cols = 2
max_rows = 10
table = doc.tables[1]
prev_text = ""

print("Assessment: ")

### Assessment
### Original loop from link below
### https://www.reddit.com/r/learnpython/comments/dbie6s/help_iterating_over_a_table_in_pythondocx/
for row in table.rows:
    if (row_iter == 0):
        row_iter += 1
        continue
    data.append(("SLO " + row.cells[first_col].text, row.cells[sec_col].text, row.cells[third_col].text, row.cells[fourth_col].text))

print(data)

data.clear()

row_iter = 0
table = doc.tables[2]
prev_text = ""

print("\nData:")

### Data Analysis
### Original loop from link below
### https://www.reddit.com/r/learnpython/comments/dbie6s/help_iterating_over_a_table_in_pythondocx/
for row in table.rows:
    if row.cells[first_col].text == "SLO " + str(slo_count + 1) + ": ":
        break
    if row_iter <= one_cols or (row_iter > max_rows and row_iter <= one_cols + max_rows):
        data.append(row.cells[first_col].text)
    if (row_iter > one_cols and row_iter < max_rows + 1) or (row_iter > max_rows + one_cols):
        data.append({row.cells[0].text: row.cells[sec_col].text})
    row_iter += 1
    if row_iter > max_rows:
        row_iter = 0

print(data)

print(chkbox_element_list)

#### DATA COLLECTION AND ANALYSIS
data_coll_list = []
table = doc.tables[3]
for row in table.rows:
    if "SLO " + str(slo_count + 1) in row.cells[first_col].text:
        break
    slo_measure = row.cells[first_col].text
    data_coll_range = row.cells[sec_col].text
    num_students = row.cells[third_col].text
    percent_students = row.cells[fourth_col].text
    extend_values = [(slo_measure, data_coll_range, num_students, percent_students)]
    if data_coll_range == "" and num_students == "" and percent_students == "":
        continue
    data_coll_list.extend(extend_values)

print(dec_act)